
from langchain.callbacks import get_openai_callback
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers.json import SimpleJsonOutputParser
from langchain.prompts import PromptTemplate
import pathlib
import textwrap

import google.generativeai as genai

from IPython.display import display
from IPython.display import Markdown

from typing import List, Optional
from pydantic import BaseModel, Field
from docx import Document

# def to_markdown(text):
#   text = text.replace('•', '  *')
#   return Markdown(textwrap.indent(text, '> ', pre dicate=lambda _: True))

# genai.configure(api_key="AIzaSyA2ToY7E9HcedmcMSRaeN0MNaw6eGSWLnA")
# model = genai.GenerativeModel('gemini-pro')



llm = ChatOpenAI(model_name="gpt-3.5-turbo-0125", temperature=0,openai_api_key="api_key")
format_instructions = """
            {
                "name": "Candidate Name",
                "email": "Candidate Email",  
                "phone": "Candidate Phone",
                
                "education": [
                    {
                        "college": "College Name",
                        "degree": "Degree Name",
                        "start_date": "Start Date",
                        "end_date": "End Date"
                    }
                ],
                "experiences": [
                    {
                        "company": "Company Name",
                        "title": "Job Title",
                        "start_date": "Start Date",
                        "end_date": "End Date",
                        "duration": "Job Duration",
                        "location": "Location"
                    }
                ],
                
                "summary": "profile summary",
                "skills":[
                "list of skills"
                ],
                "Certifications or Honor-Awards": [
                    {
                        "certificate_title":"certificate title",
                        "certificate_provider":"certificate provider"
                    }
                ],
                "others":"other important and valuable information in CV" 
            }
        """
template = PromptTemplate(
            template="""Extract only the instructed information from the candidate's LinkedIn CV and output it in the specified JSON format.
                        Do not include any additional entries.:
                        {format_instructions}.
                        And Candidate CV:
                        {cv_text}
                        Note: Do not extrac information that is not there in the cv_test
                        """,
            input_variables=["cv_text"],
            partial_variables={
                "format_instructions": format_instructions,
            },
            validate_template=False,
)

import csv
import os
def read_pdf(file_path):
    text = ""
    with fitz.open(file_path) as pdf_document:
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text += page.get_text()
    return text
def read_pdfs_in_folder(folder_path):
    pdf_texts = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".pdf"):
            file_path = os.path.join(folder_path, file_name)
            pdf_texts.append(read_pdf(file_path))
    return pdf_texts
def read_word_files_in_folder(folder_path):
    word_texts = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            doc = Document(file_path)
            text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
            word_texts.append(text)
    return word_texts
def save_to_csv(data, csv_file_path):
    csv_exists = os.path.isfile(csv_file_path)
    with open(csv_file_path, 'a', newline='', encoding='utf-8') as csvfile:
        csv_writer = csv.DictWriter(csvfile, fieldnames=data.keys())
        if not csv_exists:
            csv_writer.writeheader()
            csv_exists = True
        csv_writer.writerow(data)
        print(f"Data appended to CSV file '{csv_file_path}' successfully.")

import time
import fitz  # PyMuPDF
folder_path = "C:/Users/Jonathan Khawbs/Desktop/Internship Folder/resumes_linkedin" # Replace with the actual path to your folder


pdf_texts = read_pdfs_in_folder(folder_path)

print(pdf_texts)
print(len(pdf_texts))
# Combine both lists of texts

def extract_text_from_files(folder_path):
    all_texts = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)  # Use os.path.join for cross-platform compatibility

        if file_name.endswith(".doc"):  # Changed from ".doc" to ".docx" to target Word documents
            try:
                doc = Document(file_path)
                text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                all_texts.append(text)
            except Exception as e:
                print(f"Error processing {file_name}: {e}")
                continue  # Skip this file and continue with the next one

    return all_texts

doc_files=extract_text_from_files(folder_path)

print(len(doc_files))
all_files=doc_files+pdf_texts
print(len(all_files))

csv_file_path = "rani_mendez.csv"  # Change the file name if needed

start_time=time.time()
text = all_files
print(text)
data1=[]
with get_openai_callback() as cb:
    for texts in text:

        try:
            st=time.time()
            output = llm.invoke(template.format(cv_text=texts))

            json_parser = SimpleJsonOutputParser()
            res = json_parser.parse(output.content)
            # response = model.generate_content(template.format(cv_text=text))
            # print(response.text)

            e=time.time()
            time_taken=e-st


            data = {"model":[],
              'time_take':time_taken,
              'Total Tokens': cb.total_tokens,
              'Prompt   Tokens': cb.prompt_tokens,
              'Completion Tokens': cb.completion_tokens,
              'Total Cost (USD)': cb.total_cost}

            data["model"].append(llm.model_name)

            save_to_csv(res, csv_file_path)

            # print(f"time_taken to save in csv= {time}")
            res.update(data)
            print(res)

        except Exception as e:
            print(f"Error processing entry: {e}")
            continue
print(cb)
end_time = time.time()
elapsed_time = end_time - start_time
print(f"elapsed time={elapsed_time}")
print(data1)

